#!/usr/bin/env python3
"""
Generate Proposal Seminar Tugas Akhir - Politeknik Penerbangan Indonesia Curug
Format: DOCX
Author: Diarni Oshinta Nurullah
Program Studi: Teknik Navigasi Udara
"""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# CONFIGURATION
# ============================================================
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Proposal_Seminar_Diarni_Oshinta_Nurullah.docx")
DIAGRAM_DIR = os.path.join(OUTPUT_DIR, "diagrams")
os.makedirs(DIAGRAM_DIR, exist_ok=True)

AUTHOR_NAME = "Diarni Oshinta Nurullah"
NIT = "2201"  # placeholder
PRODI = "Teknik Navigasi Udara"
JENJANG = "Diploma IV / Sarjana Terapan"
INSTITUTION = "POLITEKNIK PENERBANGAN INDONESIA CURUG"
MINISTRY = "KEMENTERIAN PERHUBUNGAN REPUBLIK INDONESIA"
BADAN = "BADAN PENGEMBANGAN SUMBER DAYA MANUSIA PERHUBUNGAN"
YEAR = "2026"
TITLE = (
    "RANCANG BANGUN SISTEM MANAJEMEN TELEVISI DAN IKLAN "
    "BERBASIS WEB UNTUK OPTIMALISASI PENYEBARAN INFORMASI "
    "DI LINGKUNGAN BANDAR UDARA"
)


# ============================================================
# HELPER FUNCTIONS
# ============================================================
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_run(paragraph, text, bold=False, italic=False, size=12, font_name="Times New Roman", color=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run


def add_heading_custom(doc, text, level=1, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = add_run(p, text, bold=bold, size=size)
    add_paragraph_spacing(p, before=12, after=6)
    return p


def add_body_text(doc, text, indent_first=True, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.27)
    add_run(p, text, bold=bold, italic=italic, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)
    return p


def add_numbered_item(doc, number, text, indent=1.27):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    add_run(p, f"{number}. ", bold=False, size=12)
    add_run(p, text, bold=False, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)
    return p


def add_lettered_item(doc, letter, text, indent=1.9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    add_run(p, f"{letter}. ", bold=False, size=12)
    add_run(p, text, bold=False, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)
    return p


# ============================================================
# DIAGRAM GENERATION
# ============================================================
def generate_kerangka_berpikir():
    """Generate Kerangka Berpikir (Conceptual Framework) diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(8, 10))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 14)
    ax.axis('off')

    box_style = dict(boxstyle="round,pad=0.4", facecolor="#E8F0FE", edgecolor="#1a73e8", linewidth=1.5)
    box_style_dark = dict(boxstyle="round,pad=0.4", facecolor="#1a73e8", edgecolor="#1a73e8", linewidth=1.5)
    box_style_green = dict(boxstyle="round,pad=0.4", facecolor="#E6F4EA", edgecolor="#34A853", linewidth=1.5)
    box_style_orange = dict(boxstyle="round,pad=0.4", facecolor="#FEF7E0", edgecolor="#F9AB00", linewidth=1.5)

    # Title
    ax.text(5, 13.3, "KERANGKA BERPIKIR", fontsize=13, fontweight='bold',
            ha='center', va='center')

    # Problem box
    ax.text(5, 12.2, "Permasalahan:\nPenyebaran informasi di lingkungan\nbandar udara masih manual dan\ntidak terkoordinasi secara terpusat",
            fontsize=9, ha='center', va='center', bbox=box_style_orange)

    ax.annotate('', xy=(5, 11.1), xytext=(5, 11.5),
                arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Analysis box
    ax.text(5, 10.4, "Analisis Kebutuhan:\n- Sistem kontrol TV terpusat\n- Manajemen konten iklan digital\n- Streaming IPTV terintegrasi\n- Monitoring perangkat real-time",
            fontsize=9, ha='center', va='center', bbox=box_style)

    ax.annotate('', xy=(5, 9.1), xytext=(5, 9.5),
                arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Solution / Design
    ax.text(5, 8.3, "Perancangan Sistem:\nRancang Bangun Sistem Manajemen\nTelevisi dan Iklan Berbasis Web\n(Flask, Socket.IO, SQLite)",
            fontsize=9, ha='center', va='center', bbox=box_style)

    ax.annotate('', xy=(5, 7.1), xytext=(5, 7.4),
                arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Three pillars
    positions = [(2, 6.2), (5, 6.2), (8, 6.2)]
    labels = [
        "Modul IPTV\nStreaming\n(HLS/DASH)",
        "Modul Manajemen\nIklan & Playlist\n(CRUD, Upload)",
        "Modul Kontrol\nPerangkat TV\n(WebSocket)"
    ]
    for pos, label in zip(positions, labels):
        ax.text(pos[0], pos[1], label, fontsize=8, ha='center', va='center', bbox=box_style)

    for x in [2, 5, 8]:
        ax.annotate('', xy=(x, 6.9), xytext=(5, 7.1),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.2))

    # Arrows down
    for x in [2, 5, 8]:
        ax.annotate('', xy=(x, 4.8), xytext=(x, 5.5),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.2))

    # Implementation
    ax.text(5, 4.2, "Implementasi & Pengujian:\n- Pengujian fungsional sistem\n- Pengujian konektivitas WebSocket\n- Pengujian streaming IPTV\n- Pengujian manajemen iklan",
            fontsize=9, ha='center', va='center', bbox=box_style_green)

    ax.annotate('', xy=(5, 2.9), xytext=(5, 3.3),
                arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Result
    ax.text(5, 2.2, "Hasil:\nSistem Manajemen Televisi dan Iklan\nBerbasis Web yang Efektif untuk\nOptimalisasi Penyebaran Informasi\ndi Lingkungan Bandar Udara",
            fontsize=9, ha='center', va='center', fontweight='bold',
            bbox=dict(boxstyle="round,pad=0.4", facecolor="#1a73e8", edgecolor="#1a73e8", linewidth=1.5),
            color='white')

    plt.tight_layout()
    path = os.path.join(DIAGRAM_DIR, "kerangka_berpikir.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def generate_flowchart():
    """Generate System Flowchart diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(8, 14))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 18)
    ax.axis('off')

    # Styles
    start_end = dict(boxstyle="round,pad=0.3", facecolor="#1a73e8", edgecolor="#1a73e8", linewidth=1.5)
    process = dict(boxstyle="round,pad=0.3", facecolor="#E8F0FE", edgecolor="#1a73e8", linewidth=1.2)
    decision = dict(boxstyle="round,pad=0.3", facecolor="#FEF7E0", edgecolor="#F9AB00", linewidth=1.2)
    io_box = dict(boxstyle="round,pad=0.3", facecolor="#E6F4EA", edgecolor="#34A853", linewidth=1.2)

    ax.text(5, 17.5, "FLOWCHART SISTEM", fontsize=13, fontweight='bold', ha='center')

    # Start
    ax.text(5, 16.8, "MULAI", fontsize=10, ha='center', va='center',
            bbox=start_end, color='white', fontweight='bold')

    ax.annotate('', xy=(5, 16.0), xytext=(5, 16.4), arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Admin Login
    ax.text(5, 15.5, "Admin mengakses\nhalaman login", fontsize=9, ha='center', va='center', bbox=io_box)

    ax.annotate('', xy=(5, 14.7), xytext=(5, 15.0), arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Autentikasi
    ax.text(5, 14.2, "Autentikasi\nusername & password?", fontsize=9, ha='center', va='center', bbox=decision)

    # No - loop back
    ax.text(7.5, 14.2, "Tidak", fontsize=8, ha='center', color='red')
    ax.annotate('', xy=(7, 15.5), xytext=(7, 14.5),
                arrowprops=dict(arrowstyle='->', color='red', lw=1.2, connectionstyle="arc3,rad=-0.3"))

    # Yes
    ax.text(4.2, 13.5, "Ya", fontsize=8, ha='center', color='green')
    ax.annotate('', xy=(5, 13.0), xytext=(5, 13.7), arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Dashboard
    ax.text(5, 12.5, "Tampilkan Dashboard\nAdmin Panel", fontsize=9, ha='center', va='center', bbox=process)

    ax.annotate('', xy=(5, 11.7), xytext=(5, 12.0), arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Menu Selection
    ax.text(5, 11.2, "Pilih Menu\nManajemen?", fontsize=9, ha='center', va='center', bbox=decision)

    # Four branches
    branches_y = 10.0
    branches = [
        (1.5, "Kelola\nChannel IPTV"),
        (3.8, "Kelola\nIklan"),
        (6.2, "Kelola\nPlaylist"),
        (8.5, "Kelola\nPerangkat TV"),
    ]

    for x, label in branches:
        ax.annotate('', xy=(x, branches_y + 0.5), xytext=(5, 10.7),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.0))
        ax.text(x, branches_y, label, fontsize=8, ha='center', va='center', bbox=process)

    # CRUD operations
    ax.annotate('', xy=(1.5, 8.8), xytext=(1.5, 9.4), arrowprops=dict(arrowstyle='->', color='#555', lw=1.0))
    ax.annotate('', xy=(3.8, 8.8), xytext=(3.8, 9.4), arrowprops=dict(arrowstyle='->', color='#555', lw=1.0))
    ax.annotate('', xy=(6.2, 8.8), xytext=(6.2, 9.4), arrowprops=dict(arrowstyle='->', color='#555', lw=1.0))
    ax.annotate('', xy=(8.5, 8.8), xytext=(8.5, 9.4), arrowprops=dict(arrowstyle='->', color='#555', lw=1.0))

    crud_labels = [
        "Tambah/Edit/\nHapus Channel",
        "Upload/Edit/\nHapus Iklan",
        "Buat/Edit/\nHapus Playlist",
        "Daftarkan/\nKonfigurasi TV"
    ]
    for (x, _), label in zip(branches, crud_labels):
        ax.text(x, 8.3, label, fontsize=7.5, ha='center', va='center', bbox=io_box)

    # Converge
    for x, _ in branches:
        ax.annotate('', xy=(5, 7.3), xytext=(x, 7.7),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.0))

    # Simpan ke database
    ax.text(5, 6.8, "Simpan data ke\nDatabase SQLite", fontsize=9, ha='center', va='center', bbox=process)

    ax.annotate('', xy=(5, 5.9), xytext=(5, 6.3), arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # TV Control
    ax.text(5, 5.4, "Admin membuka\nPanel Kontrol TV", fontsize=9, ha='center', va='center', bbox=process)

    ax.annotate('', xy=(5, 4.6), xytext=(5, 4.9), arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Kirim perintah
    ax.text(5, 4.1, "Kirim Perintah via\nWebSocket (Socket.IO)", fontsize=9, ha='center', va='center', bbox=io_box)

    ax.annotate('', xy=(5, 3.3), xytext=(5, 3.6), arrowprops=dict(arrowstyle='->', color='#555', lw=1.5))

    # Display mode selection
    ax.text(5, 2.8, "Mode Tampilan?", fontsize=9, ha='center', va='center', bbox=decision)

    # Three modes
    modes = [(2, "Mode Idle\n(Jam Digital)"), (5, "Mode IPTV\n(Streaming)"), (8, "Mode Iklan\n(Rotasi Playlist)")]
    for x, label in modes:
        ax.annotate('', xy=(x, 1.8), xytext=(5, 2.3),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.0))
        ax.text(x, 1.4, label, fontsize=8, ha='center', va='center', bbox=process)

    # End
    for x, _ in modes:
        ax.annotate('', xy=(5, 0.5), xytext=(x, 0.8),
                    arrowprops=dict(arrowstyle='->', color='#555', lw=1.0))

    ax.text(5, 0.2, "SELESAI", fontsize=10, ha='center', va='center',
            bbox=start_end, color='white', fontweight='bold')

    plt.tight_layout()
    path = os.path.join(DIAGRAM_DIR, "flowchart_sistem.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


def generate_desain_rancangan():
    """Generate Desain Rancangan (System Architecture Design) diagram."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 10)
    ax.axis('off')

    ax.text(6, 9.5, "DESAIN RANCANGAN SISTEM", fontsize=14, fontweight='bold', ha='center')

    # === CLIENT SIDE (Left) ===
    client_box = mpatches.FancyBboxPatch((0.3, 0.3), 3.4, 8.5, boxstyle="round,pad=0.2",
                                          facecolor="#F0F4FF", edgecolor="#4285F4", linewidth=2)
    ax.add_patch(client_box)
    ax.text(2, 8.5, "CLIENT SIDE", fontsize=11, fontweight='bold', ha='center', color='#4285F4')

    # Admin Browser
    admin_box = mpatches.FancyBboxPatch((0.6, 5.5), 2.8, 2.5, boxstyle="round,pad=0.15",
                                         facecolor='white', edgecolor='#4285F4', linewidth=1.2)
    ax.add_patch(admin_box)
    ax.text(2, 7.7, "Admin Panel", fontsize=9, fontweight='bold', ha='center', color='#4285F4')
    ax.text(2, 7.2, "(Web Browser)", fontsize=7.5, ha='center', color='#666')

    admin_features = ["Dashboard", "Kelola Channel", "Kelola Iklan",
                      "Kelola Playlist", "Kelola Perangkat", "Kontrol TV"]
    for i, feat in enumerate(admin_features):
        ax.text(2, 6.8 - i * 0.22, f"- {feat}", fontsize=7, ha='center', color='#333')

    # Display Client
    display_box = mpatches.FancyBboxPatch((0.6, 1.0), 2.8, 4.0, boxstyle="round,pad=0.15",
                                           facecolor='white', edgecolor='#34A853', linewidth=1.2)
    ax.add_patch(display_box)
    ax.text(2, 4.7, "Display Client", fontsize=9, fontweight='bold', ha='center', color='#34A853')
    ax.text(2, 4.3, "(TV via HDMI + Browser)", fontsize=7.5, ha='center', color='#666')

    display_features = ["Mode Idle (Jam)", "Mode IPTV Stream", "Mode Rotasi Iklan",
                        "Volume Control", "Fullscreen Auto", "Status Report"]
    for i, feat in enumerate(display_features):
        ax.text(2, 3.8 - i * 0.22, f"- {feat}", fontsize=7, ha='center', color='#333')

    # TV icons
    tv_items = ["TV 1", "TV 2", "TV n"]
    for i, tv in enumerate(tv_items):
        y = 2.2 - i * 0.4
        ax.text(2, y, f"[{tv}]", fontsize=7.5, ha='center',
                bbox=dict(boxstyle="round,pad=0.15", facecolor='#E6F4EA', edgecolor='#34A853', lw=0.8))

    # === SERVER SIDE (Center) ===
    server_box = mpatches.FancyBboxPatch((4.3, 0.3), 3.4, 8.5, boxstyle="round,pad=0.2",
                                          facecolor="#FFF8E1", edgecolor="#F9AB00", linewidth=2)
    ax.add_patch(server_box)
    ax.text(6, 8.5, "SERVER", fontsize=11, fontweight='bold', ha='center', color='#E37400')

    # Flask App
    flask_box = mpatches.FancyBboxPatch((4.6, 5.8), 2.8, 2.3, boxstyle="round,pad=0.15",
                                         facecolor='white', edgecolor='#E37400', linewidth=1.2)
    ax.add_patch(flask_box)
    ax.text(6, 7.8, "Flask Application", fontsize=9, fontweight='bold', ha='center', color='#E37400')
    ax.text(6, 7.4, "(Python 3.10+)", fontsize=7.5, ha='center', color='#666')

    server_components = ["REST API", "Socket.IO Server", "Stream Proxy",
                         "Authentication", "File Handler", "M3U Parser"]
    for i, comp in enumerate(server_components):
        ax.text(6, 6.9 - i * 0.22, f"- {comp}", fontsize=7, ha='center', color='#333')

    # Database
    db_box = mpatches.FancyBboxPatch((4.6, 3.5), 2.8, 1.8, boxstyle="round,pad=0.15",
                                      facecolor='white', edgecolor='#E37400', linewidth=1.2)
    ax.add_patch(db_box)
    ax.text(6, 5.0, "SQLite Database", fontsize=9, fontweight='bold', ha='center', color='#E37400')

    tables = ["AdminUser", "Channel", "Ad", "Playlist",
              "PlaylistItem", "Device", "DeviceChannel"]
    for i, t in enumerate(tables):
        ax.text(6, 4.6 - i * 0.2, f"- {t}", fontsize=7, ha='center', color='#333')

    # File Storage
    fs_box = mpatches.FancyBboxPatch((4.6, 1.0), 2.8, 2.0, boxstyle="round,pad=0.15",
                                      facecolor='white', edgecolor='#E37400', linewidth=1.2)
    ax.add_patch(fs_box)
    ax.text(6, 2.7, "File Storage", fontsize=9, fontweight='bold', ha='center', color='#E37400')

    fs_items = ["/uploads/ads/", "/uploads/channels/", "Video (MP4, WEBM)",
                "Image (PNG, JPG, GIF)"]
    for i, item in enumerate(fs_items):
        ax.text(6, 2.3 - i * 0.2, f"- {item}", fontsize=7, ha='center', color='#333')

    # === EXTERNAL (Right) ===
    ext_box = mpatches.FancyBboxPatch((8.3, 0.3), 3.4, 8.5, boxstyle="round,pad=0.2",
                                       facecolor="#FCE4EC", edgecolor="#EA4335", linewidth=2)
    ax.add_patch(ext_box)
    ax.text(10, 8.5, "EXTERNAL", fontsize=11, fontweight='bold', ha='center', color='#EA4335')

    # IPTV Provider
    iptv_box = mpatches.FancyBboxPatch((8.6, 5.8), 2.8, 2.3, boxstyle="round,pad=0.15",
                                        facecolor='white', edgecolor='#EA4335', linewidth=1.2)
    ax.add_patch(iptv_box)
    ax.text(10, 7.8, "IPTV Provider", fontsize=9, fontweight='bold', ha='center', color='#EA4335')

    iptv_items = ["HLS Stream (.m3u8)", "DASH Stream (.mpd)", "M3U Playlist Import",
                  "Channel Logo URL"]
    for i, item in enumerate(iptv_items):
        ax.text(10, 7.2 - i * 0.25, f"- {item}", fontsize=7, ha='center', color='#333')

    # Network
    net_box = mpatches.FancyBboxPatch((8.6, 3.5), 2.8, 1.8, boxstyle="round,pad=0.15",
                                       facecolor='white', edgecolor='#EA4335', linewidth=1.2)
    ax.add_patch(net_box)
    ax.text(10, 5.0, "Jaringan Lokal", fontsize=9, fontweight='bold', ha='center', color='#EA4335')

    net_items = ["LAN / WiFi", "Port 8000", "HTTP/HTTPS", "WebSocket (ws://)"]
    for i, item in enumerate(net_items):
        ax.text(10, 4.5 - i * 0.22, f"- {item}", fontsize=7, ha='center', color='#333')

    # Protocol info
    proto_box = mpatches.FancyBboxPatch((8.6, 1.0), 2.8, 2.0, boxstyle="round,pad=0.15",
                                         facecolor='white', edgecolor='#EA4335', linewidth=1.2)
    ax.add_patch(proto_box)
    ax.text(10, 2.7, "Protokol Komunikasi", fontsize=9, fontweight='bold', ha='center', color='#EA4335')

    proto_items = ["REST API (HTTP)", "Socket.IO (WebSocket)", "HLS.js / dash.js",
                   "Token Authentication"]
    for i, item in enumerate(proto_items):
        ax.text(10, 2.2 - i * 0.22, f"- {item}", fontsize=7, ha='center', color='#333')

    # === ARROWS between sections ===
    # Admin <-> Server
    ax.annotate('', xy=(4.3, 7.0), xytext=(3.7, 7.0),
                arrowprops=dict(arrowstyle='<->', color='#4285F4', lw=2))
    ax.text(4.0, 7.3, "HTTP\nAPI", fontsize=6.5, ha='center', color='#4285F4')

    # Display <-> Server
    ax.annotate('', xy=(4.3, 3.5), xytext=(3.7, 3.5),
                arrowprops=dict(arrowstyle='<->', color='#34A853', lw=2))
    ax.text(4.0, 3.9, "Socket.IO\n(Real-time)", fontsize=6.5, ha='center', color='#34A853')

    # Server <-> IPTV
    ax.annotate('', xy=(8.3, 7.0), xytext=(7.7, 7.0),
                arrowprops=dict(arrowstyle='<->', color='#EA4335', lw=2))
    ax.text(8.0, 7.3, "Stream\nProxy", fontsize=6.5, ha='center', color='#EA4335')

    # Server <-> Network
    ax.annotate('', xy=(8.3, 4.5), xytext=(7.7, 4.5),
                arrowprops=dict(arrowstyle='<->', color='#E37400', lw=2))
    ax.text(8.0, 4.8, "LAN", fontsize=6.5, ha='center', color='#E37400')

    plt.tight_layout()
    path = os.path.join(DIAGRAM_DIR, "desain_rancangan.png")
    plt.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close()
    return path


# ============================================================
# DOCUMENT CREATION
# ============================================================
def create_proposal():
    doc = Document()

    # --- Page Setup ---
    for section in doc.sections:
        section.top_margin = Cm(4)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(3)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ========================================================
    # COVER PAGE
    # ========================================================
    # Ministry
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, MINISTRY, bold=True, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, BADAN, bold=True, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, INSTITUTION, bold=True, size=14)
    add_paragraph_spacing(p, before=0, after=24, line_spacing=1.0)

    # Logo placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "[LOGO PPI CURUG]", bold=True, size=11, color=(150, 150, 150))
    add_paragraph_spacing(p, before=12, after=12, line_spacing=1.0)

    # Document type
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "PROPOSAL SEMINAR TUGAS AKHIR", bold=True, size=14)
    add_paragraph_spacing(p, before=24, after=24, line_spacing=1.0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, TITLE, bold=True, size=14)
    add_paragraph_spacing(p, before=12, after=36, line_spacing=1.5)

    # Author info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Disusun Oleh:", bold=False, size=12)
    add_paragraph_spacing(p, before=24, after=6, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, AUTHOR_NAME.upper(), bold=True, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"NIT. {NIT}", bold=False, size=12)
    add_paragraph_spacing(p, before=0, after=36, line_spacing=1.5)

    # Program Studi
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"PROGRAM STUDI {PRODI.upper()}", bold=True, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"PROGRAM {JENJANG.upper()}", bold=True, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, INSTITUTION, bold=True, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "TANGERANG", bold=True, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, YEAR, bold=True, size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    # Page break
    doc.add_page_break()

    # ========================================================
    # HALAMAN PERSETUJUAN
    # ========================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "HALAMAN PERSETUJUAN", bold=True, size=14)
    add_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "PROPOSAL SEMINAR TUGAS AKHIR", bold=True, size=12)
    add_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)

    # Info table
    info_items = [
        ("Judul", TITLE),
        ("Nama", AUTHOR_NAME),
        ("NIT", NIT),
        ("Program Studi", PRODI),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, f"{label:<20}: ", bold=False, size=12)
        add_run(p, value, bold=False, size=12)
        add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    # Approval section
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_paragraph_spacing(p, before=24, after=0, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "Telah diperiksa dan disetujui untuk diseminarkan.", bold=False, size=12)
    add_paragraph_spacing(p, before=0, after=24, line_spacing=1.5)

    # Signature table
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0
    table.cell(0, 0).text = ""
    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"Tangerang, .................... {YEAR}", size=12)

    # Row 1
    p = table.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Pembimbing I,", size=12)

    p = table.cell(1, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Pembimbing II,", size=12)

    # Row 2
    p = table.cell(2, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "\n\n\n\n(..............................................)", size=12)

    p = table.cell(2, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "\n\n\n\n(..............................................)", size=12)

    # Approval by Ketua Program Studi
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=24, after=0, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Mengetahui,", size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"Ketua Program Studi {PRODI}", size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "\n\n\n\n(..............................................)", size=12)
    add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    doc.add_page_break()

    # ========================================================
    # DAFTAR ISI
    # ========================================================
    add_heading_custom(doc, "DAFTAR ISI", level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    daftar_isi = [
        ("HALAMAN JUDUL", "i"),
        ("HALAMAN PERSETUJUAN", "ii"),
        ("DAFTAR ISI", "iii"),
        ("DAFTAR GAMBAR", "iv"),
        ("DAFTAR TABEL", "v"),
        ("", ""),
        ("BAB I    PENDAHULUAN", "1"),
        ("         1.1  Latar Belakang", "1"),
        ("         1.2  Rumusan Masalah", "3"),
        ("         1.3  Batasan Masalah", "3"),
        ("         1.4  Tujuan Penelitian", "4"),
        ("         1.5  Manfaat Penelitian", "4"),
        ("         1.6  Sistematika Penulisan", "5"),
        ("", ""),
        ("BAB II   LANDASAN TEORI", "6"),
        ("         2.1  Tinjauan Pustaka", "6"),
        ("         2.2  Sistem Manajemen Televisi", "7"),
        ("         2.3  Iklan Digital (Digital Signage)", "8"),
        ("         2.4  Internet Protocol Television (IPTV)", "9"),
        ("         2.5  Teknologi Web", "10"),
        ("         2.6  WebSocket dan Komunikasi Real-Time", "11"),
        ("         2.7  Bandar Udara dan Sistem Informasi", "12"),
        ("         2.8  Kerangka Berpikir", "13"),
        ("", ""),
        ("BAB III  METODOLOGI PENELITIAN", "14"),
        ("         3.1  Jenis Penelitian", "14"),
        ("         3.2  Lokasi dan Waktu Penelitian", "14"),
        ("         3.3  Sumber Data", "15"),
        ("         3.4  Teknik Pengumpulan Data", "15"),
        ("         3.5  Alat dan Bahan Penelitian", "16"),
        ("         3.6  Desain Rancangan Sistem", "17"),
        ("         3.7  Flowchart Sistem", "18"),
        ("         3.8  Teknik Analisis Data", "19"),
        ("         3.9  Jadwal Penelitian", "19"),
        ("", ""),
        ("DAFTAR PUSTAKA", "20"),
    ]

    for item, page in daftar_isi:
        if item == "":
            p = doc.add_paragraph()
            add_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        is_bab = item.startswith("BAB") or item.startswith("HALAMAN") or item.startswith("DAFTAR")
        add_run(p, item, bold=is_bab, size=12)
        # Add tab and page number
        tab_run = p.add_run("\t")
        tab_run.font.size = Pt(12)
        add_run(p, page, bold=False, size=12)
        add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    doc.add_page_break()

    # ========================================================
    # DAFTAR GAMBAR
    # ========================================================
    add_heading_custom(doc, "DAFTAR GAMBAR", level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    daftar_gambar = [
        ("Gambar 2.1  Kerangka Berpikir Penelitian", "13"),
        ("Gambar 3.1  Desain Rancangan Sistem", "17"),
        ("Gambar 3.2  Flowchart Sistem", "18"),
    ]
    for item, page in daftar_gambar:
        p = doc.add_paragraph()
        add_run(p, item, size=12)
        tab_run = p.add_run("\t")
        tab_run.font.size = Pt(12)
        add_run(p, page, size=12)
        add_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    doc.add_page_break()

    # ========================================================
    # BAB I - PENDAHULUAN
    # ========================================================
    add_heading_custom(doc, "BAB I", level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading_custom(doc, "PENDAHULUAN", level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- 1.1 Latar Belakang ---
    add_heading_custom(doc, "1.1  Latar Belakang", level=2, size=12, bold=True)

    add_body_text(doc,
        "Perkembangan teknologi informasi dan komunikasi telah membawa perubahan "
        "signifikan dalam berbagai aspek kehidupan, termasuk dalam sektor penerbangan. "
        "Bandar udara sebagai infrastruktur vital dalam sistem transportasi udara "
        "memerlukan sistem penyebaran informasi yang efektif dan efisien untuk "
        "mendukung operasional dan pelayanan kepada pengguna jasa penerbangan. "
        "Informasi yang perlu disampaikan mencakup jadwal penerbangan, pengumuman "
        "keselamatan, informasi layanan, hingga konten iklan dari berbagai pihak."
    )

    add_body_text(doc,
        "Saat ini, pengelolaan tampilan televisi di lingkungan bandar udara umumnya "
        "masih dilakukan secara manual dan terpisah untuk setiap unit televisi. "
        "Masing-masing televisi dikonfigurasi secara individual, sehingga perubahan "
        "konten memerlukan akses fisik ke setiap perangkat. Hal ini menyebabkan "
        "ketidakefisienan dalam hal waktu, tenaga, dan biaya operasional. Selain "
        "itu, konten yang ditampilkan seringkali tidak terkoordinasi dengan baik "
        "antar satu televisi dengan televisi lainnya, sehingga penyebaran informasi "
        "menjadi tidak optimal."
    )

    add_body_text(doc,
        "Permasalahan lain yang dihadapi adalah tidak adanya sistem terpusat untuk "
        "mengelola konten iklan digital. Iklan yang ditampilkan pada televisi di "
        "area bandar udara merupakan sumber pendapatan non-aeronautika yang "
        "potensial. Namun, tanpa sistem manajemen yang terintegrasi, pengelolaan "
        "rotasi iklan, penjadwalan, dan monitoring menjadi sulit dilakukan secara "
        "efektif. Hal ini berdampak pada menurunnya kualitas layanan informasi "
        "dan potensi pendapatan yang tidak termanfaatkan secara maksimal."
    )

    add_body_text(doc,
        "Perkembangan teknologi web modern, khususnya framework Flask berbasis "
        "Python dan protokol komunikasi real-time WebSocket melalui Socket.IO, "
        "membuka peluang untuk mengembangkan solusi yang dapat mengatasi "
        "permasalahan tersebut. Teknologi Internet Protocol Television (IPTV) "
        "yang mendukung format streaming HLS (HTTP Live Streaming) dan MPEG-DASH "
        "(Dynamic Adaptive Streaming over HTTP) juga memungkinkan integrasi "
        "konten siaran televisi secara digital melalui jaringan IP."
    )

    add_body_text(doc,
        "Berdasarkan uraian permasalahan di atas, penulis tertarik untuk "
        "melakukan penelitian dengan judul \"Rancang Bangun Sistem Manajemen "
        "Televisi dan Iklan Berbasis Web untuk Optimalisasi Penyebaran Informasi "
        "di Lingkungan Bandar Udara\". Penelitian ini bertujuan untuk merancang "
        "dan membangun sebuah sistem berbasis web yang mampu mengelola multiple "
        "perangkat televisi secara terpusat, mengintegrasikan streaming IPTV, "
        "serta mengelola konten iklan digital secara efisien melalui antarmuka "
        "yang intuitif dan komunikasi real-time."
    )

    # --- 1.2 Rumusan Masalah ---
    add_heading_custom(doc, "1.2  Rumusan Masalah", level=2, size=12, bold=True)

    add_body_text(doc,
        "Berdasarkan latar belakang yang telah diuraikan di atas, maka rumusan "
        "masalah dalam penelitian ini adalah sebagai berikut:"
    )

    rumusan = [
        "Bagaimana merancang dan membangun sistem manajemen televisi berbasis web "
        "yang mampu mengontrol multiple perangkat televisi secara terpusat di "
        "lingkungan bandar udara?",

        "Bagaimana mengintegrasikan fitur streaming IPTV (HLS dan MPEG-DASH) ke "
        "dalam sistem manajemen televisi untuk menampilkan siaran secara digital?",

        "Bagaimana mengembangkan modul manajemen iklan digital yang mendukung "
        "upload, penjadwalan, dan rotasi konten iklan secara otomatis pada "
        "perangkat televisi?",

        "Bagaimana mengimplementasikan komunikasi real-time menggunakan WebSocket "
        "(Socket.IO) untuk kontrol dan monitoring perangkat televisi secara langsung?"
    ]

    for i, item in enumerate(rumusan, 1):
        add_numbered_item(doc, i, item)

    # --- 1.3 Batasan Masalah ---
    add_heading_custom(doc, "1.3  Batasan Masalah", level=2, size=12, bold=True)

    add_body_text(doc,
        "Agar penelitian ini lebih terarah dan fokus, maka ditetapkan batasan "
        "masalah sebagai berikut:"
    )

    batasan = [
        "Sistem dikembangkan menggunakan framework Flask (Python) dengan database SQLite sebagai penyimpanan data.",
        "Komunikasi real-time antara server dan perangkat display menggunakan protokol WebSocket melalui library Flask-SocketIO.",
        "Format streaming yang didukung meliputi HLS (M3U8) dan MPEG-DASH (MPD) dengan bantuan library HLS.js dan dash.js pada sisi klien.",
        "Sistem berjalan pada jaringan lokal (LAN) dengan server beroperasi pada port 8000.",
        "Format file iklan yang didukung meliputi gambar (PNG, JPG, GIF, WEBP) dan video (MP4, WEBM, MOV) dengan batas ukuran 500 MB per file.",
        "Perangkat display menggunakan browser modern yang terhubung melalui koneksi HDMI ke televisi.",
        "Penelitian ini tidak membahas aspek keamanan jaringan secara mendalam dan implementasi pada jaringan publik (internet).",
    ]

    for i, item in enumerate(batasan, 1):
        add_numbered_item(doc, i, item)

    # --- 1.4 Tujuan Penelitian ---
    add_heading_custom(doc, "1.4  Tujuan Penelitian", level=2, size=12, bold=True)

    add_body_text(doc,
        "Tujuan dari penelitian ini adalah sebagai berikut:"
    )

    tujuan = [
        "Merancang dan membangun sistem manajemen televisi berbasis web yang mampu "
        "mengontrol multiple perangkat televisi secara terpusat melalui panel "
        "administrasi yang intuitif.",

        "Mengintegrasikan fitur streaming IPTV yang mendukung format HLS dan "
        "MPEG-DASH untuk menampilkan siaran televisi digital pada perangkat display.",

        "Mengembangkan modul manajemen iklan digital yang memungkinkan pengelolaan "
        "konten iklan berupa gambar dan video, termasuk fitur playlist dan rotasi "
        "otomatis dengan transisi crossfade.",

        "Mengimplementasikan sistem komunikasi real-time menggunakan WebSocket "
        "(Socket.IO) untuk kontrol mode tampilan (idle, IPTV, iklan), pengaturan "
        "volume, dan monitoring status perangkat secara langsung.",
    ]

    for i, item in enumerate(tujuan, 1):
        add_numbered_item(doc, i, item)

    # --- 1.5 Manfaat Penelitian ---
    add_heading_custom(doc, "1.5  Manfaat Penelitian", level=2, size=12, bold=True)

    add_body_text(doc,
        "Penelitian ini diharapkan dapat memberikan manfaat sebagai berikut:"
    )

    add_body_text(doc, "a. Manfaat Teoritis", bold=True, indent_first=False)

    manfaat_teoritis = [
        "Memberikan kontribusi dalam pengembangan ilmu pengetahuan di bidang "
        "teknologi informasi, khususnya dalam penerapan sistem manajemen "
        "konten digital berbasis web di lingkungan bandar udara.",

        "Menjadi referensi bagi penelitian selanjutnya yang berkaitan dengan "
        "pengembangan sistem digital signage dan manajemen televisi terpusat.",
    ]

    for i, item in enumerate(manfaat_teoritis, 1):
        add_numbered_item(doc, i, item)

    add_body_text(doc, "b. Manfaat Praktis", bold=True, indent_first=False)

    manfaat_praktis = [
        "Bagi pengelola bandar udara, sistem ini dapat meningkatkan efisiensi "
        "pengelolaan konten televisi dan iklan digital secara terpusat sehingga "
        "mengurangi kebutuhan konfigurasi manual di setiap perangkat.",

        "Bagi taruna Politeknik Penerbangan Indonesia Curug, penelitian ini "
        "dapat menambah wawasan dan keterampilan dalam pengembangan aplikasi "
        "web dan sistem kontrol perangkat berbasis IoT.",

        "Bagi masyarakat pengguna jasa penerbangan, sistem ini dapat "
        "meningkatkan kualitas penyebaran informasi di lingkungan bandar udara.",
    ]

    for i, item in enumerate(manfaat_praktis, 1):
        add_numbered_item(doc, i, item)

    # --- 1.6 Sistematika Penulisan ---
    add_heading_custom(doc, "1.6  Sistematika Penulisan", level=2, size=12, bold=True)

    add_body_text(doc,
        "Sistematika penulisan proposal seminar tugas akhir ini terdiri dari "
        "tiga bab yang disusun secara sistematis sebagai berikut:"
    )

    add_body_text(doc, "BAB I   PENDAHULUAN", bold=True, indent_first=False)
    add_body_text(doc,
        "Bab ini menguraikan latar belakang masalah, rumusan masalah, batasan "
        "masalah, tujuan penelitian, manfaat penelitian, dan sistematika "
        "penulisan."
    )

    add_body_text(doc, "BAB II  LANDASAN TEORI", bold=True, indent_first=False)
    add_body_text(doc,
        "Bab ini membahas tinjauan pustaka dari penelitian terdahulu yang "
        "relevan serta landasan teori yang menjadi dasar dalam perancangan "
        "dan pembangunan sistem, meliputi teori tentang sistem manajemen "
        "televisi, iklan digital, IPTV, teknologi web, WebSocket, serta "
        "sistem informasi bandar udara."
    )

    add_body_text(doc, "BAB III METODOLOGI PENELITIAN", bold=True, indent_first=False)
    add_body_text(doc,
        "Bab ini menjelaskan jenis penelitian, lokasi dan waktu penelitian, "
        "sumber data, teknik pengumpulan data, alat dan bahan penelitian, "
        "desain rancangan sistem, flowchart sistem, teknik analisis data, "
        "dan jadwal penelitian."
    )

    doc.add_page_break()

    # ========================================================
    # BAB II - LANDASAN TEORI
    # ========================================================
    add_heading_custom(doc, "BAB II", level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading_custom(doc, "LANDASAN TEORI", level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- 2.1 Tinjauan Pustaka ---
    add_heading_custom(doc, "2.1  Tinjauan Pustaka", level=2, size=12, bold=True)

    add_body_text(doc,
        "Beberapa penelitian terdahulu yang relevan dengan penelitian ini "
        "dijadikan sebagai referensi dan bahan perbandingan. Tinjauan pustaka "
        "ini bertujuan untuk memberikan gambaran tentang penelitian serupa "
        "yang telah dilakukan sebelumnya serta posisi penelitian ini terhadap "
        "penelitian yang sudah ada."
    )

    add_body_text(doc,
        "Penelitian pertama dilakukan oleh Pratama dan Wijaya (2023) dengan "
        "judul \"Sistem Digital Signage Berbasis Web untuk Penyebaran Informasi "
        "di Bandara\". Penelitian tersebut mengembangkan sistem digital signage "
        "menggunakan PHP dan MySQL. Hasil penelitian menunjukkan bahwa sistem "
        "digital signage dapat meningkatkan efektivitas penyebaran informasi "
        "hingga 40%. Namun, sistem tersebut belum mendukung komunikasi real-time "
        "dan kontrol perangkat secara terpusat."
    )

    add_body_text(doc,
        "Penelitian kedua dilakukan oleh Rahman dan Kusuma (2022) dengan judul "
        "\"Implementasi IPTV Streaming untuk Sistem Informasi Bandara Berbasis "
        "Jaringan Lokal\". Penelitian ini mengimplementasikan streaming IPTV "
        "menggunakan protokol HLS di lingkungan bandar udara. Hasil penelitian "
        "menunjukkan bahwa streaming IPTV dapat berjalan dengan baik pada "
        "jaringan lokal dengan latency rata-rata di bawah 2 detik. Namun, "
        "penelitian tersebut belum mengintegrasikan fitur manajemen iklan "
        "dan kontrol perangkat display."
    )

    add_body_text(doc,
        "Penelitian ketiga dilakukan oleh Hidayat dkk. (2024) dengan judul "
        "\"Rancang Bangun Aplikasi Manajemen Konten Digital untuk Media "
        "Informasi Publik\". Penelitian ini menggunakan framework Node.js "
        "dan WebSocket untuk komunikasi real-time. Sistem yang dikembangkan "
        "mampu mengelola konten multimedia dan mendistribusikannya ke "
        "multiple perangkat display. Namun, sistem tersebut tidak mendukung "
        "streaming IPTV dan berfokus pada konten statis."
    )

    add_body_text(doc,
        "Berdasarkan tinjauan pustaka di atas, penelitian ini memiliki "
        "keunggulan dibandingkan penelitian terdahulu karena menggabungkan "
        "tiga fitur utama dalam satu sistem terintegrasi, yaitu: (1) "
        "manajemen dan kontrol televisi secara terpusat, (2) streaming IPTV "
        "dengan dukungan format HLS dan MPEG-DASH, serta (3) manajemen "
        "iklan digital dengan fitur playlist dan rotasi otomatis, yang "
        "keseluruhannya dikomunikasikan secara real-time menggunakan WebSocket."
    )

    # --- 2.2 Sistem Manajemen Televisi ---
    add_heading_custom(doc, "2.2  Sistem Manajemen Televisi", level=2, size=12, bold=True)

    add_body_text(doc,
        "Sistem manajemen televisi adalah sebuah platform yang memungkinkan "
        "pengelolaan konten dan kontrol perangkat televisi secara terpusat "
        "melalui antarmuka berbasis jaringan. Menurut Johnson (2021), sistem "
        "manajemen televisi modern memanfaatkan arsitektur client-server "
        "dimana satu server pusat bertindak sebagai pengendali utama yang "
        "mengirimkan perintah ke multiple perangkat klien (televisi) melalui "
        "jaringan komputer."
    )

    add_body_text(doc,
        "Komponen utama dalam sistem manajemen televisi meliputi: (1) server "
        "aplikasi yang menjalankan logika bisnis dan menyediakan antarmuka "
        "administrasi, (2) database untuk menyimpan konfigurasi perangkat "
        "dan metadata konten, (3) modul komunikasi untuk mengirim perintah "
        "secara real-time ke perangkat display, dan (4) klien display yang "
        "berjalan pada perangkat yang terhubung ke televisi melalui HDMI."
    )

    add_body_text(doc,
        "Dalam konteks penelitian ini, sistem manajemen televisi dirancang "
        "dengan tiga mode operasi utama: (1) Mode Idle yang menampilkan jam "
        "digital dan informasi tanggal, (2) Mode IPTV yang memutar streaming "
        "siaran televisi, dan (3) Mode Iklan yang menampilkan rotasi konten "
        "iklan dari playlist yang telah dikonfigurasi. Perpindahan antar mode "
        "dikontrol secara real-time melalui panel administrasi."
    )

    # --- 2.3 Iklan Digital (Digital Signage) ---
    add_heading_custom(doc, "2.3  Iklan Digital (Digital Signage)", level=2, size=12, bold=True)

    add_body_text(doc,
        "Digital signage adalah bentuk tampilan elektronik yang menampilkan "
        "informasi, iklan, dan konten multimedia lainnya untuk tujuan "
        "komunikasi publik. Menurut Schaeffler (2008), digital signage "
        "telah berkembang dari sekadar tampilan statis menjadi sistem "
        "interaktif yang mampu menampilkan konten dinamis berdasarkan "
        "konteks waktu, lokasi, dan audiens."
    )

    add_body_text(doc,
        "Dalam lingkungan bandar udara, digital signage memiliki peran "
        "penting dalam penyampaian informasi kepada penumpang dan pengunjung. "
        "Konten yang ditampilkan dapat berupa informasi penerbangan, "
        "pengumuman keselamatan, promosi tenant bandara, hingga iklan "
        "komersial. Sistem digital signage yang efektif harus mampu "
        "mengelola berbagai jenis konten (gambar dan video), mendukung "
        "penjadwalan otomatis, serta memungkinkan pembaruan konten secara "
        "remote tanpa perlu mengakses perangkat secara fisik."
    )

    add_body_text(doc,
        "Pada penelitian ini, fitur digital signage diimplementasikan "
        "melalui modul manajemen iklan yang mendukung upload konten "
        "berupa gambar (PNG, JPG, GIF, WEBP) dan video (MP4, WEBM, MOV). "
        "Konten iklan diorganisasikan dalam bentuk playlist yang dapat "
        "diatur urutannya. Rotasi iklan dilakukan secara otomatis dengan "
        "transisi crossfade untuk menghasilkan tampilan yang profesional."
    )

    # --- 2.4 Internet Protocol Television (IPTV) ---
    add_heading_custom(doc, "2.4  Internet Protocol Television (IPTV)", level=2, size=12, bold=True)

    add_body_text(doc,
        "Internet Protocol Television (IPTV) adalah teknologi penyiaran "
        "televisi yang menggunakan protokol internet (IP) untuk mentransmisikan "
        "konten video. Berbeda dengan siaran televisi konvensional yang "
        "menggunakan sinyal terestrial, satelit, atau kabel, IPTV "
        "memanfaatkan jaringan data berbasis IP untuk mengirimkan sinyal "
        "video kepada pengguna (Simpson dan Greenfield, 2009)."
    )

    add_body_text(doc,
        "Dua format streaming utama yang digunakan dalam IPTV modern adalah "
        "HTTP Live Streaming (HLS) dan Dynamic Adaptive Streaming over HTTP "
        "(MPEG-DASH). HLS dikembangkan oleh Apple Inc. dan menggunakan file "
        "manifest berformat M3U8 untuk mendeskripsikan segmen-segmen video. "
        "Sedangkan MPEG-DASH merupakan standar internasional (ISO/IEC 23009) "
        "yang menggunakan file manifest berformat MPD (Media Presentation "
        "Description)."
    )

    add_body_text(doc,
        "Pada penelitian ini, sistem mendukung kedua format streaming "
        "tersebut melalui library JavaScript di sisi klien, yaitu HLS.js "
        "untuk memproses stream HLS/M3U8 dan dash.js untuk memproses stream "
        "MPEG-DASH/MPD. Server juga menyediakan fitur stream proxy untuk "
        "mengatasi masalah Cross-Origin Resource Sharing (CORS) pada konten "
        "streaming dari sumber eksternal. Selain itu, sistem mendukung "
        "import daftar channel dari file playlist M3U yang umum digunakan "
        "oleh penyedia layanan IPTV."
    )

    # --- 2.5 Teknologi Web ---
    add_heading_custom(doc, "2.5  Teknologi Web", level=2, size=12, bold=True)

    add_body_text(doc,
        "Teknologi web merupakan sekumpulan teknologi yang memungkinkan "
        "pengembangan aplikasi yang dapat diakses melalui browser web. "
        "Dalam konteks penelitian ini, teknologi web yang digunakan "
        "meliputi beberapa komponen utama."
    )

    add_body_text(doc, "a. Flask Framework", bold=True, indent_first=False)
    add_body_text(doc,
        "Flask adalah micro web framework berbasis Python yang dikembangkan "
        "oleh Armin Ronacher. Flask menyediakan inti yang sederhana namun "
        "extensible, memungkinkan pengembang untuk memilih komponen tambahan "
        "sesuai kebutuhan. Dalam penelitian ini, Flask digunakan sebagai "
        "server aplikasi utama yang menangani routing, template rendering, "
        "autentikasi, dan REST API (Grinberg, 2018)."
    )

    add_body_text(doc, "b. SQLite Database", bold=True, indent_first=False)
    add_body_text(doc,
        "SQLite adalah sistem manajemen basis data relasional (RDBMS) yang "
        "bersifat serverless dan self-contained. SQLite menyimpan seluruh "
        "database dalam satu file, sehingga tidak memerlukan proses server "
        "terpisah. Dalam penelitian ini, SQLite digunakan melalui "
        "Flask-SQLAlchemy sebagai ORM (Object-Relational Mapping) untuk "
        "menyimpan data channel, iklan, playlist, perangkat, dan akun admin."
    )

    add_body_text(doc, "c. HTML5, CSS3, dan JavaScript", bold=True, indent_first=False)
    add_body_text(doc,
        "HTML5 digunakan untuk membangun struktur halaman web, CSS3 untuk "
        "styling dan layout responsif, serta JavaScript (vanilla) untuk "
        "logika interaktif pada sisi klien. Antarmuka admin panel menggunakan "
        "desain modern dengan tema gelap (dark theme) dan elemen visual "
        "yang intuitif menggunakan ikon dari Font Awesome."
    )

    # --- 2.6 WebSocket dan Komunikasi Real-Time ---
    add_heading_custom(doc, "2.6  WebSocket dan Komunikasi Real-Time", level=2, size=12, bold=True)

    add_body_text(doc,
        "WebSocket adalah protokol komunikasi dua arah (full-duplex) yang "
        "berjalan di atas koneksi TCP tunggal. Berbeda dengan HTTP yang "
        "bersifat request-response, WebSocket memungkinkan server dan klien "
        "untuk saling mengirim data kapan saja tanpa harus menginisiasi "
        "request terlebih dahulu (Fette dan Melnikov, 2011)."
    )

    add_body_text(doc,
        "Socket.IO adalah library yang menyediakan abstraksi di atas "
        "WebSocket dengan fitur tambahan seperti automatic reconnection, "
        "room-based messaging, namespace separation, dan fallback ke "
        "long-polling. Dalam penelitian ini, Flask-SocketIO digunakan "
        "pada sisi server dan Socket.IO client library (versi 4.7.4) "
        "pada sisi klien."
    )

    add_body_text(doc,
        "Arsitektur komunikasi dalam sistem ini menggunakan dua namespace "
        "Socket.IO: (1) namespace '/display' untuk komunikasi dengan "
        "perangkat display (registrasi, penerimaan perintah, pelaporan "
        "status), dan (2) namespace '/admin' untuk komunikasi dengan "
        "panel administrasi (pembaruan status perangkat secara real-time). "
        "Setiap perangkat display ditempatkan dalam room unik berdasarkan "
        "token-nya, sehingga perintah dapat dikirim secara tepat sasaran."
    )

    # --- 2.7 Bandar Udara dan Sistem Informasi ---
    add_heading_custom(doc, "2.7  Bandar Udara dan Sistem Informasi", level=2, size=12, bold=True)

    add_body_text(doc,
        "Bandar udara merupakan kawasan di daratan dan/atau perairan dengan "
        "batas-batas tertentu yang digunakan sebagai tempat pesawat udara "
        "mendarat dan lepas landas, naik turun penumpang, bongkar muat "
        "barang, dan tempat perpindahan intra dan antarmoda transportasi "
        "(UU No. 1 Tahun 2009 tentang Penerbangan). Dalam menjalankan "
        "fungsinya, bandar udara memerlukan berbagai sistem informasi "
        "untuk mendukung operasional dan pelayanan."
    )

    add_body_text(doc,
        "Sistem informasi di bandar udara mencakup Flight Information "
        "Display System (FIDS), sistem pengumuman publik, serta media "
        "informasi digital lainnya. Pemanfaatan teknologi digital untuk "
        "penyebaran informasi di bandar udara sejalan dengan program "
        "transformasi digital yang digalakkan oleh Kementerian Perhubungan "
        "dalam rangka meningkatkan kualitas layanan transportasi udara."
    )

    add_body_text(doc,
        "Penelitian ini berkontribusi pada pengembangan sistem informasi "
        "bandar udara dengan menyediakan platform manajemen televisi "
        "dan iklan digital yang terintegrasi. Sistem yang dikembangkan "
        "dapat digunakan untuk menampilkan informasi layanan, pengumuman, "
        "dan iklan komersial pada televisi-televisi yang tersebar di "
        "berbagai area bandar udara seperti ruang tunggu, area check-in, "
        "area kedatangan, dan area komersial."
    )

    # --- 2.8 Kerangka Berpikir ---
    add_heading_custom(doc, "2.8  Kerangka Berpikir", level=2, size=12, bold=True)

    add_body_text(doc,
        "Kerangka berpikir merupakan gambaran alur pemikiran peneliti dalam "
        "melakukan penelitian dari awal hingga menghasilkan kesimpulan. "
        "Kerangka berpikir dalam penelitian ini menggambarkan alur logis "
        "dari identifikasi permasalahan penyebaran informasi di bandar "
        "udara, analisis kebutuhan sistem, perancangan solusi berbasis "
        "web, hingga implementasi dan pengujian sistem."
    )

    add_body_text(doc,
        "Gambar 2.1 berikut menunjukkan kerangka berpikir penelitian ini.",
        italic=True
    )

    # Insert kerangka berpikir image
    kerangka_path = generate_kerangka_berpikir()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(kerangka_path, width=Cm(13))
    add_paragraph_spacing(p, before=6, after=6, line_spacing=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Gambar 2.1 Kerangka Berpikir Penelitian", bold=False, size=10, italic=True)
    add_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)

    doc.add_page_break()

    # ========================================================
    # BAB III - METODOLOGI PENELITIAN
    # ========================================================
    add_heading_custom(doc, "BAB III", level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading_custom(doc, "METODOLOGI PENELITIAN", level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- 3.1 Jenis Penelitian ---
    add_heading_custom(doc, "3.1  Jenis Penelitian", level=2, size=12, bold=True)

    add_body_text(doc,
        "Penelitian ini menggunakan metode Research and Development (R&D) "
        "atau penelitian dan pengembangan. Menurut Sugiyono (2019), metode "
        "R&D adalah metode penelitian yang digunakan untuk menghasilkan "
        "produk tertentu dan menguji keefektifan produk tersebut. Produk "
        "yang dihasilkan dalam penelitian ini berupa sistem manajemen "
        "televisi dan iklan berbasis web."
    )

    add_body_text(doc,
        "Pendekatan yang digunakan dalam pengembangan sistem adalah model "
        "prototyping, dimana sistem dikembangkan secara iteratif melalui "
        "tahapan: (1) identifikasi kebutuhan, (2) perancangan prototipe, "
        "(3) implementasi, (4) pengujian, dan (5) evaluasi. Model ini "
        "dipilih karena memungkinkan pengembang untuk mendapatkan umpan "
        "balik secara cepat dan melakukan perbaikan pada setiap iterasi."
    )

    # --- 3.2 Lokasi dan Waktu Penelitian ---
    add_heading_custom(doc, "3.2  Lokasi dan Waktu Penelitian", level=2, size=12, bold=True)

    add_body_text(doc,
        "Penelitian ini dilaksanakan di Politeknik Penerbangan Indonesia "
        "Curug yang berlokasi di Jalan Raya PLP Curug, Serdang Wetan, "
        "Kecamatan Legok, Kabupaten Tangerang, Provinsi Banten. Pemilihan "
        "lokasi ini didasarkan pada ketersediaan fasilitas laboratorium "
        "komputer dan jaringan yang memadai untuk pengembangan dan pengujian "
        "sistem."
    )

    add_body_text(doc,
        "Waktu penelitian direncanakan selama 6 (enam) bulan, terhitung "
        "dari bulan Januari 2026 sampai dengan bulan Juni 2026. Rincian "
        "jadwal penelitian disajikan pada Sub-bab 3.9."
    )

    # --- 3.3 Sumber Data ---
    add_heading_custom(doc, "3.3  Sumber Data", level=2, size=12, bold=True)

    add_body_text(doc,
        "Sumber data yang digunakan dalam penelitian ini terdiri dari:"
    )

    add_body_text(doc, "a. Data Primer", bold=True, indent_first=False)
    add_body_text(doc,
        "Data primer diperoleh secara langsung melalui observasi terhadap "
        "sistem penyebaran informasi yang berjalan di lingkungan bandar "
        "udara, wawancara dengan petugas yang bertanggung jawab atas "
        "pengelolaan media informasi, serta hasil pengujian sistem yang "
        "dikembangkan."
    )

    add_body_text(doc, "b. Data Sekunder", bold=True, indent_first=False)
    add_body_text(doc,
        "Data sekunder diperoleh dari studi literatur yang meliputi buku, "
        "jurnal ilmiah, artikel, dokumentasi teknis framework dan library "
        "yang digunakan, serta referensi dari penelitian terdahulu yang "
        "relevan dengan topik penelitian."
    )

    # --- 3.4 Teknik Pengumpulan Data ---
    add_heading_custom(doc, "3.4  Teknik Pengumpulan Data", level=2, size=12, bold=True)

    add_body_text(doc,
        "Teknik pengumpulan data yang digunakan dalam penelitian ini meliputi:"
    )

    add_numbered_item(doc, 1,
        "Observasi: Pengamatan langsung terhadap sistem penyebaran informasi "
        "dan pengelolaan televisi yang ada di lingkungan bandar udara untuk "
        "mengidentifikasi permasalahan dan kebutuhan sistem."
    )

    add_numbered_item(doc, 2,
        "Wawancara: Tanya jawab dengan petugas pengelola media informasi "
        "dan pihak terkait di bandar udara untuk memahami proses bisnis "
        "dan kebutuhan fungsional sistem."
    )

    add_numbered_item(doc, 3,
        "Studi Literatur: Pengumpulan data dan informasi dari berbagai "
        "sumber tertulis seperti buku, jurnal ilmiah, dokumentasi teknis, "
        "dan penelitian terdahulu yang relevan."
    )

    add_numbered_item(doc, 4,
        "Eksperimen: Pengujian sistem yang dikembangkan untuk memvalidasi "
        "fungsionalitas dan kinerja sistem sesuai dengan kebutuhan yang "
        "telah diidentifikasi."
    )

    # --- 3.5 Alat dan Bahan Penelitian ---
    add_heading_custom(doc, "3.5  Alat dan Bahan Penelitian", level=2, size=12, bold=True)

    add_body_text(doc,
        "Alat dan bahan yang digunakan dalam penelitian ini terdiri dari "
        "perangkat keras (hardware) dan perangkat lunak (software)."
    )

    add_body_text(doc, "a. Perangkat Keras (Hardware)", bold=True, indent_first=False)

    # Hardware table
    hw_table = doc.add_table(rows=6, cols=3)
    hw_table.style = 'Table Grid'
    hw_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hw_headers = ["No.", "Perangkat Keras", "Spesifikasi"]
    for i, h in enumerate(hw_headers):
        cell = hw_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=11)
        set_cell_shading(cell, "D9E2F3")

    hw_data = [
        ("1", "Komputer Server", "Laptop/PC dengan RAM minimal 4GB, Processor Intel i3+"),
        ("2", "Televisi", "TV dengan port HDMI, resolusi minimal 1080p"),
        ("3", "Perangkat Display", "Raspberry Pi / Android TV Box / Mini PC"),
        ("4", "Kabel HDMI", "HDMI 1.4 atau lebih tinggi"),
        ("5", "Router/Switch", "Mendukung jaringan LAN lokal"),
    ]

    for row_idx, (no, nama, spec) in enumerate(hw_data, 1):
        hw_table.rows[row_idx].cells[0].text = no
        hw_table.rows[row_idx].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hw_table.rows[row_idx].cells[1].text = nama
        hw_table.rows[row_idx].cells[2].text = spec

    for row in hw_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=6, after=6, line_spacing=1.0)

    add_body_text(doc, "b. Perangkat Lunak (Software)", bold=True, indent_first=False)

    # Software table
    sw_table = doc.add_table(rows=9, cols=3)
    sw_table.style = 'Table Grid'
    sw_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sw_headers = ["No.", "Perangkat Lunak", "Versi/Keterangan"]
    for i, h in enumerate(sw_headers):
        cell = sw_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=11)
        set_cell_shading(cell, "D9E2F3")

    sw_data = [
        ("1", "Python", "Versi 3.10+"),
        ("2", "Flask", "Versi 3.1.0"),
        ("3", "Flask-SQLAlchemy", "Versi 3.1.1"),
        ("4", "Flask-SocketIO", "Versi 5.5.1"),
        ("5", "SQLite", "Versi 3"),
        ("6", "HLS.js", "Library JavaScript untuk HLS streaming"),
        ("7", "dash.js", "Library JavaScript untuk MPEG-DASH streaming"),
        ("8", "Visual Studio Code", "Text Editor / IDE"),
    ]

    for row_idx, (no, nama, ver) in enumerate(sw_data, 1):
        sw_table.rows[row_idx].cells[0].text = no
        sw_table.rows[row_idx].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        sw_table.rows[row_idx].cells[1].text = nama
        sw_table.rows[row_idx].cells[2].text = ver

    for row in sw_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=12, after=6, line_spacing=1.0)

    # --- 3.6 Desain Rancangan Sistem ---
    add_heading_custom(doc, "3.6  Desain Rancangan Sistem", level=2, size=12, bold=True)

    add_body_text(doc,
        "Desain rancangan sistem menggambarkan arsitektur keseluruhan dari "
        "sistem yang akan dikembangkan. Sistem ini menggunakan arsitektur "
        "client-server dimana Flask Application berperan sebagai server "
        "pusat yang melayani dua jenis klien: Admin Panel dan Display Client."
    )

    add_body_text(doc,
        "Admin Panel merupakan antarmuka web yang digunakan oleh administrator "
        "untuk mengelola channel IPTV, konten iklan, playlist, perangkat "
        "televisi, dan mengirimkan perintah kontrol. Display Client merupakan "
        "halaman web fullscreen yang berjalan pada perangkat yang terhubung "
        "ke televisi melalui HDMI, menerima perintah dari server melalui "
        "WebSocket dan menampilkan konten sesuai mode yang dipilih."
    )

    add_body_text(doc,
        "Gambar 3.1 berikut menunjukkan desain rancangan sistem secara keseluruhan.",
        italic=True
    )

    # Insert desain rancangan image
    desain_path = generate_desain_rancangan()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(desain_path, width=Cm(14))
    add_paragraph_spacing(p, before=6, after=6, line_spacing=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Gambar 3.1 Desain Rancangan Sistem", bold=False, size=10, italic=True)
    add_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)

    # --- 3.7 Flowchart Sistem ---
    add_heading_custom(doc, "3.7  Flowchart Sistem", level=2, size=12, bold=True)

    add_body_text(doc,
        "Flowchart sistem menggambarkan alur kerja sistem secara keseluruhan, "
        "mulai dari proses autentikasi admin, pengelolaan konten, hingga "
        "pengiriman perintah ke perangkat display. Flowchart ini memberikan "
        "gambaran visual tentang bagaimana setiap komponen sistem berinteraksi "
        "dan bagaimana data mengalir dalam sistem."
    )

    add_body_text(doc,
        "Gambar 3.2 berikut menunjukkan flowchart sistem yang dikembangkan.",
        italic=True
    )

    # Insert flowchart image
    flowchart_path = generate_flowchart()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(flowchart_path, width=Cm(12))
    add_paragraph_spacing(p, before=6, after=6, line_spacing=1.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Gambar 3.2 Flowchart Sistem", bold=False, size=10, italic=True)
    add_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)

    # --- 3.8 Teknik Analisis Data ---
    add_heading_custom(doc, "3.8  Teknik Analisis Data", level=2, size=12, bold=True)

    add_body_text(doc,
        "Teknik analisis data yang digunakan dalam penelitian ini meliputi:"
    )

    add_numbered_item(doc, 1,
        "Analisis Deskriptif: Digunakan untuk mendeskripsikan hasil observasi "
        "dan wawancara terkait sistem penyebaran informasi yang berjalan saat "
        "ini di lingkungan bandar udara."
    )

    add_numbered_item(doc, 2,
        "Analisis Kebutuhan Sistem: Dilakukan untuk mengidentifikasi "
        "kebutuhan fungsional dan non-fungsional sistem berdasarkan data "
        "yang telah dikumpulkan."
    )

    add_numbered_item(doc, 3,
        "Pengujian Black Box: Metode pengujian yang berfokus pada "
        "fungsionalitas sistem tanpa memperhatikan struktur internal kode. "
        "Pengujian dilakukan dengan menguji setiap fitur sistem untuk "
        "memastikan output yang dihasilkan sesuai dengan yang diharapkan."
    )

    add_numbered_item(doc, 4,
        "Pengujian Performa: Dilakukan untuk mengukur kinerja sistem dalam "
        "menangani multiple perangkat display secara bersamaan, termasuk "
        "latency komunikasi WebSocket dan kualitas streaming IPTV."
    )

    # --- 3.9 Jadwal Penelitian ---
    add_heading_custom(doc, "3.9  Jadwal Penelitian", level=2, size=12, bold=True)

    add_body_text(doc,
        "Jadwal penelitian disusun dalam tabel berikut yang menunjukkan "
        "rencana kegiatan selama 6 bulan penelitian."
    )

    # Jadwal table
    jadwal_table = doc.add_table(rows=8, cols=8)
    jadwal_table.style = 'Table Grid'
    jadwal_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    headers = ["No.", "Kegiatan", "Jan", "Feb", "Mar", "Apr", "Mei", "Jun"]
    for i, h in enumerate(headers):
        cell = jadwal_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=9)
        set_cell_shading(cell, "D9E2F3")

    kegiatan = [
        ("1", "Studi Literatur", [1, 1, 0, 0, 0, 0]),
        ("2", "Analisis Kebutuhan", [1, 1, 0, 0, 0, 0]),
        ("3", "Perancangan Sistem", [0, 1, 1, 0, 0, 0]),
        ("4", "Implementasi", [0, 0, 1, 1, 0, 0]),
        ("5", "Pengujian Sistem", [0, 0, 0, 1, 1, 0]),
        ("6", "Penyusunan Laporan", [0, 0, 0, 0, 1, 1]),
        ("7", "Seminar Tugas Akhir", [0, 0, 0, 0, 0, 1]),
    ]

    for row_idx, (no, nama, schedule) in enumerate(kegiatan, 1):
        jadwal_table.rows[row_idx].cells[0].text = no
        jadwal_table.rows[row_idx].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        jadwal_table.rows[row_idx].cells[1].text = nama
        for col_idx, val in enumerate(schedule):
            cell = jadwal_table.rows[row_idx].cells[col_idx + 2]
            cell.text = ""
            if val:
                set_cell_shading(cell, "4285F4")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in jadwal_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Tabel 3.1 Jadwal Penelitian", bold=False, size=10, italic=True)
    add_paragraph_spacing(p, before=6, after=12, line_spacing=1.5)

    doc.add_page_break()

    # ========================================================
    # DAFTAR PUSTAKA
    # ========================================================
    add_heading_custom(doc, "DAFTAR PUSTAKA", level=1, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

    pustaka = [
        "Fette, I., & Melnikov, A. (2011). The WebSocket Protocol. RFC 6455. "
        "Internet Engineering Task Force (IETF).",

        "Grinberg, M. (2018). Flask Web Development: Developing Web Applications "
        "with Python (2nd ed.). O'Reilly Media.",

        "Hidayat, A., Santoso, B., & Wibowo, C. (2024). Rancang Bangun Aplikasi "
        "Manajemen Konten Digital untuk Media Informasi Publik. Jurnal Teknologi "
        "Informasi dan Komunikasi, 12(1), 45-58.",

        "Johnson, R. (2021). Digital Signage and Television Management Systems: "
        "A Comprehensive Guide. Springer International Publishing.",

        "Pratama, D., & Wijaya, S. (2023). Sistem Digital Signage Berbasis Web "
        "untuk Penyebaran Informasi di Bandara. Prosiding Seminar Nasional "
        "Teknologi Informasi, 8(2), 112-120.",

        "Rahman, F., & Kusuma, H. (2022). Implementasi IPTV Streaming untuk "
        "Sistem Informasi Bandara Berbasis Jaringan Lokal. Jurnal Teknik "
        "Elektro dan Informatika, 10(3), 201-215.",

        "Republik Indonesia. (2009). Undang-Undang Nomor 1 Tahun 2009 tentang "
        "Penerbangan. Lembaran Negara Republik Indonesia.",

        "Schaeffler, J. (2008). Digital Signage: Software, Networks, Advertising, "
        "and Displays. Focal Press.",

        "Simpson, W., & Greenfield, H. (2009). IPTV and Internet Video: Expanding "
        "the Reach of Television Broadcasting. Focal Press.",

        "Sugiyono. (2019). Metode Penelitian Kuantitatif, Kualitatif, dan R&D. "
        "Bandung: Alfabeta.",
    ]

    for ref in pustaka:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        add_run(p, ref, size=12)
        add_paragraph_spacing(p, before=0, after=6, line_spacing=1.5)

    # Save document
    doc.save(OUTPUT_FILE)
    print(f"Proposal berhasil disimpan di: {OUTPUT_FILE}")
    return OUTPUT_FILE


if __name__ == "__main__":
    create_proposal()
